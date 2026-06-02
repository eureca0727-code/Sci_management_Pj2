import os
import re
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from state import ExamState, Question
import logger

_OUTPUT_DIR = "exams"

_KR_NUM = {
    "한": 1, "하나": 1, "일": 1,
    "두": 2, "둘": 2, "이": 2,
    "세": 3, "셋": 3, "삼": 3,
    "네": 4, "넷": 4, "사": 4,
    "다섯": 5, "오": 5,
    "여섯": 6, "육": 6,
    "일곱": 7, "칠": 7,
    "여덟": 8, "팔": 8,
    "아홉": 9, "구": 9,
    "열": 10, "십": 10,
}


def _parse_group_config(additional: str):
    """'N문제씩 짝지어서 M개의 대문제' 패턴 파싱 → (group_size, group_count)."""
    if not additional or "대문제" not in additional:
        return None, None

    def _to_int(token: str):
        return int(token) if token.isdigit() else _KR_NUM.get(token)

    size_m = re.search(r'(\d+|\w+)\s*문제씩', additional)
    count_m = re.search(r'(\d+|\w+)\s*개(?:의)?\s*대문제', additional)

    group_size  = _to_int(size_m.group(1))  if size_m  else 2
    group_count = _to_int(count_m.group(1)) if count_m else None
    return group_size, group_count


def _build_cover(doc: Document, requirements: dict,
                 title: str = "과학적 관리",
                 subtitle: str = "AGENT GENERATED") -> None:
    total_score = requirements.get("total_score", 100)
    today = datetime.now().strftime("%Y-%m-%d")

    # Top-right: score + date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Total  {total_score} pts    {today}")
    run.font.size = Pt(11)

    # Vertical spacer to push title toward center
    for _ in range(10):
        doc.add_paragraph()

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(36)
    run.bold = True

    # Subtitle
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor(0, 0, 0)


_META_TOKENS = re.compile(
    r"【[^】]*】|\[[^\]]*시나리오[^\]]*\]|〔[^〕]*〕"
    r"|대문제\s*\d+[^\n]*\n?|소문항\s*\d+[^\n]*\n?"
)

def _clean_content(text: str) -> str:
    """편집용 메타 토큰 제거 후 앞뒤 공백 정리."""
    return _META_TOKENS.sub("", text).strip()


def _split_app_groups(questions: list, group_config: list):
    """application 문제를 그룹(dict[gid, list])과 standalone(list)으로 분리."""
    topic_to_gid: dict[str, int] = {}
    for gc in group_config:
        for name in gc["topic_names"]:
            topic_to_gid[name] = gc["group_id"]
    grouped: dict[int, list] = {}
    standalone: list = []
    for q in questions:
        gid = topic_to_gid.get(q.get("topic", ""))
        if gid is not None:
            grouped.setdefault(gid, []).append(q)
        else:
            standalone.append(q)
    return grouped, standalone


def _build_label_map(questions: list[Question], group_config: list) -> dict[str, str]:
    """시험지 출력 순서 기준으로 {q_id: 표시 레이블} 반환.
    단독 문제 → "N."  /  그룹 소문항 → "N-1)" "N-2)"
    """
    sa    = [q for q in questions if q.get("type") == "short_answer"]
    essay = [q for q in questions if q.get("type") == "essay"]
    app   = [q for q in questions if q.get("type") == "application"]
    grouped, standalone_app = _split_app_groups(app, group_config)

    labels: dict[str, str] = {}
    seq = 1
    for q in sa:
        labels[q["id"]] = f"{seq}."
        seq += 1
    for q in essay:
        labels[q["id"]] = f"{seq}."
        seq += 1
    for gid in sorted(grouped):
        for sub_i, q in enumerate(grouped[gid], 1):
            labels[q["id"]] = f"{seq}-{sub_i})"
        seq += 1          # 그룹 전체가 하나의 번호를 소비
    for q in standalone_app:
        labels[q["id"]] = f"{seq}."
        seq += 1
    return labels


def _write_q(doc: Document, label: str, q: dict, shared_scenario: str = "") -> None:
    """단일 문제 블록 출력."""
    score = q.get("score", "?")
    scenario = _clean_content(q.get("scenario", "") or "")
    content  = _clean_content(q.get("content", ""))
    use_scenario = "" if shared_scenario else scenario

    if use_scenario:
        p = doc.add_paragraph("다음 사례를 읽고 물음에 답하시오.")
        p.runs[0].bold = True
        doc.add_paragraph(use_scenario)
        doc.add_paragraph()

    doc.add_paragraph(f"{label} ({score}pts)  {content}")
    doc.add_paragraph()


def _build_exam(doc: Document, questions: list[Question], requirements: dict,
                group_config: list | None = None,
                group_scenarios: dict | None = None):
    group_config   = group_config   or []
    group_scenarios = group_scenarios or {}

    sa_qs    = [q for q in questions if q.get("type") == "short_answer"]
    essay_qs = [q for q in questions if q.get("type") == "essay"]
    app_qs   = [q for q in questions if q.get("type") == "application"]
    grouped, standalone_app = _split_app_groups(app_qs, group_config)

    seq   = 1
    first = True  # 첫 번째 문제 앞에는 page break 없음

    # ── Part I : Short Answer ─────────────────────────────
    if sa_qs:
        _add_heading(doc, "Part I — Short Answer Questions")
        for q in sa_qs:
            if first: first = False
            else:     doc.add_page_break()
            _write_q(doc, f"{seq}.", q)
            seq += 1

    # ── Part II : Essay ───────────────────────────────────
    if essay_qs:
        if not first: doc.add_page_break()
        _add_heading(doc, "Part II — Essay Questions")
        first = False
        for i, q in enumerate(essay_qs):
            if i > 0: doc.add_page_break()
            _write_q(doc, f"{seq}.", q)
            seq += 1

    # ── Part III : Application ────────────────────────────
    if app_qs:
        if not first: doc.add_page_break()
        _add_heading(doc, "Part III — Application Questions")
        first = False
        app_item = 0   # Part III 내 출력 아이템 수

        # 그룹 문제: N. 헤딩 + 시나리오 + N-1) N-2) 소번호
        for gid in sorted(grouped):
            if app_item > 0: doc.add_page_break()
            chunk  = grouped[gid]
            shared = group_scenarios.get(gid, "")
            total  = sum(q.get("score", 0) for q in chunk)

            _add_heading(doc, f"{seq}.  ({total}pts)", level=2)

            if shared:
                p = doc.add_paragraph("다음 사례를 읽고 아래 물음에 답하시오.")
                p.runs[0].bold = True
                doc.add_paragraph(shared)
                doc.add_paragraph()

            for sub_i, q in enumerate(chunk, 1):
                if sub_i > 1: doc.add_page_break()
                _write_q(doc, f"{seq}-{sub_i})", q, shared_scenario=shared)
                app_item += 1

            seq += 1

        # 단독 application 문제
        for q in standalone_app:
            if app_item > 0: doc.add_page_break()
            _write_q(doc, f"{seq}.", q)
            seq += 1
            app_item += 1


def _build_answer_key(doc: Document, model_answers: list[dict],
                      questions: list[Question],
                      group_config: list | None = None):
    group_config = group_config or []
    label_map = _build_label_map(questions, group_config)

    # Determine display order: SA → essay → grouped app → standalone app
    sa    = [q for q in questions if q.get("type") == "short_answer"]
    essay = [q for q in questions if q.get("type") == "essay"]
    app   = [q for q in questions if q.get("type") == "application"]
    grouped, standalone_app = _split_app_groups(app, group_config)
    ordered = (sa + essay
               + [q for gid in sorted(grouped) for q in grouped[gid]]
               + standalone_app)

    answer_map = {a["question_id"]: a for a in model_answers}
    first = True
    for q in ordered:
        qid = q["id"]
        ans = answer_map.get(qid)
        if not ans:
            continue
        if first:
            first = False
        else:
            doc.add_page_break()
        label = label_map.get(qid, "?")
        _add_heading(doc, f"[{label}]  ({q.get('type','').upper()})", level=2)

        doc.add_paragraph("Model Answer:", style="Heading 3")
        doc.add_paragraph(ans.get("model_answer", ""))

        if ans.get("key_concepts"):
            doc.add_paragraph("Key Concepts: " + ", ".join(str(k) for k in ans["key_concepts"]))

        doc.add_paragraph("Rubric:", style="Heading 3")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Item", "Points", "Criteria"

        _META_ITEM = re.compile(r'\[문제\d|소문항\d|【')
        for item in ans.get("rubric", []):
            if not isinstance(item, dict):
                continue
            criteria = str(item.get("criteria", "")).strip()
            item_text = str(item.get("item", "")).strip()
            pts = item.get("points", 0)
            # 빈 행 / 0점 빈 행 / 메타 표현 행 제거
            if not criteria and not item_text:
                continue
            if _META_ITEM.search(item_text):
                continue
            try:
                if int(pts) == 0 and not criteria:
                    continue
            except (ValueError, TypeError):
                pass
            row = table.add_row().cells
            row[0].text = item_text
            row[1].text = str(pts)
            row[2].text = criteria

        doc.add_paragraph()


def run(state: ExamState) -> dict:
    logger.section("STEP 9 — Exam Compilation")
    passed        = state["passed_questions"]
    model_answers = state["model_answers"]
    requirements  = state["requirements"]

    sa_count    = sum(1 for q in passed if q.get("type") == "short_answer")
    essay_count = sum(1 for q in passed if q.get("type") == "essay")
    app_count   = sum(1 for q in passed if q.get("type") == "application")
    logger.log("Compiler", f"단답형 {sa_count} + 서술형 {essay_count} + 사례적용형 {app_count}문항 → DOCX 생성 중...")

    os.makedirs(_OUTPUT_DIR, exist_ok=True)
    ts = datetime.now().strftime("%m%d_%H%M")

    exam_doc = Document()
    _build_cover(exam_doc, requirements, subtitle="AGENT GENERATED EXAM")
    _build_exam(exam_doc, passed, requirements,
                state.get("group_config", []),
                state.get("group_scenarios", {}))
    exam_path = os.path.join(_OUTPUT_DIR, f"exam_{ts}.docx")
    exam_doc.save(exam_path)
    logger.log("Compiler", f"✅ 시험지 저장: {exam_path}")

    answer_doc = Document()
    _build_cover(answer_doc, requirements,
                 title="과학적 관리",
                 subtitle="정답 및 채점 기준")
    _build_answer_key(answer_doc, model_answers, passed,
                      state.get("group_config", []))
    answer_path = os.path.join(_OUTPUT_DIR, f"answer_key_{ts}.docx")
    answer_doc.save(answer_path)
    logger.log("Compiler", f"✅ 모범답안 저장: {answer_path}")

    return {"output_path": exam_path}
